# backend/ai_tools.py
from __future__ import annotations

import io
import logging
import os
import re
import sqlite3
import uuid
from datetime import datetime, date
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse

from .auth import AuthenticatedUser, get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ai", tags=["ai"])

# ------------------------------------------------------------
# In-memory store for generated analyses
# (If uvicorn restarts, these are lost; that's OK for v4)
# ------------------------------------------------------------
ANALYSES: Dict[str, Dict[str, Any]] = {}

# Where is the SQLite DB?
REPO_ROOT = Path(__file__).resolve().parents[1]
AO_DB_PATH = REPO_ROOT / "ao.db"

# ------------------------------------------------------------
# Small utilities
# ------------------------------------------------------------

def _utc_iso() -> str:
    return datetime.utcnow().isoformat(timespec="microseconds") + "Z"


def _normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def _safe_filename(name: str) -> str:
    name = (name or "report").strip()
    name = re.sub(r"[^\w\-.() \[\]]+", "", name, flags=re.UNICODE).strip()
    return name or "report"


def _strip_non_textual_noise(s: str) -> str:
    # helps reduce garbage when extracted from tables
    s = (s or "").replace("\x00", " ")
    s = re.sub(r"[ \t]+\n", "\n", s)
    return s


def _clip(s: str, max_len: int = 1200) -> str:
    s = _normalize_spaces(s)
    if len(s) <= max_len:
        return s
    return s[: max_len - 3].rstrip() + "..."


def _find_best_snippet(text: str, anchors: List[str], window_before: int = 300, window_after: int = 1200) -> Optional[str]:
    """
    Returns a "best effort" snippet around the first matching anchor.
    We keep it short and readable for reports.
    """
    low = (text or "").lower()
    for a in anchors:
        idx = low.find(a.lower())
        if idx != -1:
            w = text[max(0, idx - window_before): min(len(text), idx + window_after)]
            return _clip(w, 1400)
    return None


def _extract_emails(text: str) -> List[str]:
    emails = re.findall(r"(?i)\b[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}\b", text or "")
    out, seen = [], set()
    for e in emails:
        ee = e.lower()
        if ee not in seen:
            out.append(e)
            seen.add(ee)
    return out[:10]


def _extract_phones(text: str) -> List[str]:
    # very permissive: handles North American formats
    phones = re.findall(r"(?i)\b(?:\+?1[\s\-\.])?(?:\(?\d{3}\)?[\s\-\.])\d{3}[\s\-\.]\d{4}\b", text or "")
    out, seen = [], set()
    for p in phones:
        pp = _normalize_spaces(p)
        if pp not in seen:
            out.append(pp)
            seen.add(pp)
    return out[:10]


# ------------------------------------------------------------
# PDF TEXT EXTRACTION (multi-engines)
# ------------------------------------------------------------

def _extract_text_pymupdf(pdf_bytes: bytes) -> str:
    import fitz  # PyMuPDF
    text_parts: List[str] = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page in doc:
            text_parts.append(page.get_text("text") or "")
    finally:
        doc.close()
    return "\n".join(text_parts).strip()


def _extract_text_pdfplumber(pdf_bytes: bytes) -> str:
    import pdfplumber
    from io import BytesIO

    text_parts: List[str] = []
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for p in pdf.pages:
            text_parts.append(p.extract_text() or "")
    return "\n".join(text_parts).strip()


def _extract_text_pypdf(pdf_bytes: bytes) -> str:
    from io import BytesIO
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(pdf_bytes))
    text_parts: List[str] = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    Try extraction engines in order:
      1) PyMuPDF (fitz)
      2) pdfplumber
      3) pypdf (fallback)
    """
    errors: List[str] = []
    for fn in (_extract_text_pymupdf, _extract_text_pdfplumber, _extract_text_pypdf):
        try:
            txt = fn(pdf_bytes)
            txt = _strip_non_textual_noise(txt)
            if txt and len(txt.strip()) > 50:
                return txt
        except Exception as e:
            errors.append(f"{fn.__name__}: {e}")

    raise HTTPException(
        status_code=422,
        detail=f"Impossible d'extraire le texte du PDF. Détails: {errors[:3]}",
    )


# ------------------------------------------------------------
# Date parsing helpers (FR/EN)
# ------------------------------------------------------------

_FR_MONTHS = {
    "janvier": "01",
    "février": "02", "fevrier": "02",
    "mars": "03",
    "avril": "04",
    "mai": "05",
    "juin": "06",
    "juillet": "07",
    "août": "08", "aout": "08",
    "septembre": "09",
    "octobre": "10",
    "novembre": "11",
    "décembre": "12", "decembre": "12",
}

_EN_MONTHS = {
    "january": "01",
    "february": "02",
    "march": "03",
    "april": "04",
    "may": "05",
    "june": "06",
    "july": "07",
    "august": "08",
    "september": "09",
    "october": "10",
    "november": "11",
    "december": "12",
}

def _date_in_reasonable_range(iso: str) -> bool:
    try:
        y = int(iso[:4])
        return 2010 <= y <= 2035
    except Exception:
        return False


def _to_iso(y: int, m: int, d: int) -> Optional[str]:
    try:
        date(y, m, d)
        return f"{y:04d}-{m:02d}-{d:02d}"
    except Exception:
        return None


def _find_date_candidates(text: str) -> List[str]:
    out: List[str] = []
    t = text or ""

    # yyyy-mm-dd
    for m in re.finditer(r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b", t):
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        iso = _to_iso(y, mo, d)
        if iso:
            out.append(iso)

    # dd/mm/yyyy or dd-mm-yyyy
    for m in re.finditer(r"\b(\d{1,2})[/-](\d{1,2})[/-](20\d{2})\b", t):
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        iso = _to_iso(y, mo, d)
        if iso:
            out.append(iso)

    low = t.lower()

    # FR: "2 février 2024"
    for m in re.finditer(r"\b(\d{1,2})\s+([a-zéèêàùâîôûç]+)\s+(20\d{2})\b", low):
        d = int(m.group(1))
        mon = m.group(2)
        y = int(m.group(3))
        if mon in _FR_MONTHS:
            iso = _to_iso(y, int(_FR_MONTHS[mon]), d)
            if iso:
                out.append(iso)

    # EN: "february 2, 2024" or "february 2 2024"
    for m in re.finditer(r"\b([a-z]+)\s+(\d{1,2})(?:,)?\s+(20\d{2})\b", low):
        mon = m.group(1)
        d = int(m.group(2))
        y = int(m.group(3))
        if mon in _EN_MONTHS:
            iso = _to_iso(y, int(_EN_MONTHS[mon]), d)
            if iso:
                out.append(iso)

    # de-dup while preserving order
    seen = set()
    uniq = []
    for x in out:
        if x not in seen:
            seen.add(x)
            uniq.append(x)
    return uniq


def _pick_date_by_anchors(text: str, anchors: List[str]) -> Optional[str]:
    low = (text or "").lower()
    for a in anchors:
        idx = low.find(a)
        if idx != -1:
            window = low[max(0, idx - 250): min(len(low), idx + 1200)]
            cands = _find_date_candidates(window)
            cands = [d for d in cands if _date_in_reasonable_range(d)]
            if cands:
                return cands[0]
    return None


def _pick_closing_date(text: str) -> Optional[str]:
    anchors = [
        "date de clôture", "date cloture", "clôture", "cloture",
        "date limite", "date et heure limite", "soumissions doivent être reçues",
        "closing date", "closing time", "tenders must be received", "deadline",
    ]
    dt = _pick_date_by_anchors(text, anchors)
    if dt:
        return dt

    cands = _find_date_candidates((text or "")[:20000])
    cands = [d for d in cands if _date_in_reasonable_range(d)]
    return cands[0] if cands else None


def extract_key_dates(text: str) -> Dict[str, Optional[str]]:
    return {
        "closing_date": _pick_closing_date(text),
        "questions_deadline": _pick_date_by_anchors(text, [
            "date limite de questions", "date limite des questions", "questions jusqu", "questions au plus tard",
            "question deadline", "questions must be received by", "questions doivent être reçues",
        ]),
        "site_visit_date": _pick_date_by_anchors(text, [
            "visite", "visite des lieux", "visite obligatoire", "réunion d'information", "reunion d'information",
            "site visit", "mandatory site visit",
        ]),
        "addenda_deadline": _pick_date_by_anchors(text, [
            "addenda", "date limite addenda", "dernier addenda", "end of addendum",
        ]),
        "opening_date": _pick_date_by_anchors(text, [
            "ouverture des soumissions", "opening of tenders", "ouverture publique",
        ]),
    }


# ------------------------------------------------------------
# Buyer and estimated value heuristics (anti-false-positives)
# ------------------------------------------------------------

BUYER_LINE_PATTERNS = [
    r"\b(minist[eè]re\s+de\s+[^.\n]{3,120})",
    r"\b(ville\s+de\s+[^.\n]{3,120})",
    r"\b(centre\s+de\s+services\s+scolaire\s+[^.\n]{3,120})",
    r"\b(commission\s+scolaire\s+[^.\n]{3,120})",
    r"\b(agence\s+[^.\n]{3,120})",
    r"\b(universit[eé]\s+[^.\n]{3,120})",
    r"\b(cisss\s+[^.\n]{3,120})",
    r"\b(ciuss\s+[^.\n]{3,120})",
]

BAD_BUYER_SNIPPETS = {
    "complémentaire (additionnel)",
    "complementaire (additionnel)",
    "essentiel (obligatoire)",
    "obligatoire",
    "additionnel",
    "ministeres-organismes/",
    "cybersecurite-",
}


def _pick_buyer(text: str) -> Optional[str]:
    low = (text or "").lower()

    anchors = ["organisme", "client", "propriétaire", "donneur d'ouvrage", "acheteur", "buyer"]
    for a in anchors:
        idx = low.find(a)
        if idx != -1:
            window = (text or "")[max(0, idx - 200): min(len(text), idx + 800)]
            for line in window.splitlines():
                l = _normalize_spaces(line)
                ll = l.lower()
                if len(l) < 6 or len(l) > 140:
                    continue
                if any(b in ll for b in BAD_BUYER_SNIPPETS):
                    continue
                if re.search(r"https?://", ll):
                    continue
                if re.search(r"(minist[eè]re|ville|centre|commission|gouvernement|universit[eé]|agence|direction)", ll):
                    return l

    sample = (text or "")[:40000]
    low_sample = sample.lower()

    for pat in BUYER_LINE_PATTERNS:
        m = re.search(pat, low_sample, flags=re.IGNORECASE)
        if m:
            cand = _normalize_spaces(m.group(1))
            if cand and len(cand) <= 140 and cand.lower() not in BAD_BUYER_SNIPPETS:
                return cand

    return None


def _pick_estimated_value(text: str) -> Optional[str]:
    low = (text or "").lower()
    anchors = [
        "valeur estim", "valeur approximative", "budget", "montant", "plafond",
        "estimated value", "budgetary", "maximum amount", "ceiling",
    ]
    for a in anchors:
        idx = low.find(a)
        if idx != -1:
            window = (text or "")[max(0, idx - 200): min(len(text), idx + 1200)]
            money = re.findall(
                r"(?i)\b(\$?\s*\d[\d\s.,]{1,18}\s*(?:\$|cad|c\$|dollars)?)\b",
                window,
            )
            for m in money:
                s = _normalize_spaces(m)
                if not re.search(r"(?i)\$|cad|c\$|dollar", s):
                    continue
                num = re.sub(r"(?i)[^\d.,]", "", s).replace(" ", "")
                digits = re.sub(r"[^\d]", "", num)
                if not digits:
                    continue
                try:
                    val = int(digits)
                except Exception:
                    continue
                if val < 500 or val > 5_000_000_000:
                    continue
                return f"{val} CAD"

    return None


# ------------------------------------------------------------
# Extract lists in sections: mandatory req / deliverables / criteria / docs
# ------------------------------------------------------------

def _extract_list_under_heading(text: str, headings: List[str], max_lines: int = 120) -> List[str]:
    lines = [l.rstrip() for l in (text or "").splitlines()]
    low = [l.lower().strip() for l in lines]

    start = -1
    for i, ll in enumerate(low):
        if any(h in ll for h in headings):
            start = i
            break
    if start == -1:
        return []

    items: List[str] = []
    for j in range(start + 1, min(len(lines), start + 1 + max_lines)):
        raw = lines[j].strip()
        if not raw:
            if items:
                break
            continue

        # Stop if another big heading starts
        if re.match(r"^[A-ZÉÈÀÙÂÊÎÔÛ0-9][A-ZÉÈÀÙÂÊÎÔÛ\s\-]{10,}$", raw):
            break

        # bullet / numbering
        if re.match(r"^(\-|\•|\*|\d+[\.\)])\s+", raw):
            items.append(_normalize_spaces(re.sub(r"^(\-|\•|\*|\d+[\.\)])\s+", "", raw))[:380])
        elif items and len(raw) < 220:
            items[-1] = _normalize_spaces(items[-1] + " " + raw)[:520]
        else:
            if 35 <= len(raw) <= 240 and any(k in raw.lower() for k in ["%", "points", "pond", "poids", "crit", "livr", "document", "attestation", "formulaire"]):
                items.append(_normalize_spaces(raw)[:380])

    out, seen = [], set()
    for x in items:
        k = x.lower()
        if k not in seen:
            out.append(x)
            seen.add(k)
    return out[:40]


def extract_deliverables(text: str) -> List[str]:
    return _extract_list_under_heading(
        text,
        ["livrables", "produits livrables", "deliverables", "rapports", "documentation", "matériel livré", "materiel livre"],
    )


def extract_evaluation_criteria(text: str) -> List[str]:
    crit = _extract_list_under_heading(
        text,
        ["critères d’évaluation", "criteres d'evaluation", "critères d'évaluation", "évaluation", "evaluation", "grille d’évaluation", "grille d'évaluation", "pondération", "ponderation", "weighting"],
    )
    pct = [c for c in crit if re.search(r"\b\d{1,3}\s*%|\bpond|\bpoids|\bpoints\b", c.lower())]
    rest = [c for c in crit if c not in pct]
    return (pct + rest)[:40]


def extract_submission_documents(text: str) -> List[str]:
    return _extract_list_under_heading(
        text,
        [
            "documents à fournir", "documents a fournir", "documents requis", "pièces à fournir", "pieces a fournir",
            "contenu de la soumission", "contenu de l’offre", "contenu de l'offre",
            "submission must include", "submission content", "forms", "attachments", "annexes",
        ],
    )


def extract_mandatory_requirements(text: str) -> List[str]:
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    out: List[str] = []

    for i, l in enumerate(lines):
        low = l.lower()
        if "essentiel" in low and "oblig" in low:
            prev = lines[i - 1] if i - 1 >= 0 else ""
            prev2 = lines[i - 2] if i - 2 >= 0 else ""
            cand = prev
            if len(prev) < 50 and prev2 and len(prev2) < 250:
                cand = prev2 + " " + prev
            cand = _normalize_spaces(cand)
            if cand and 20 <= len(cand) <= 450:
                out.append(cand[:450])

    if not out:
        for l in lines:
            low = l.lower()
            if ("obligatoire" in low or "must" in low or "shall" in low) and ("doit" in low or "must" in low or "shall" in low):
                if 30 <= len(l) <= 280:
                    out.append(_normalize_spaces(l)[:450])

    uniq, seen = [], set()
    for x in out:
        k = x.lower()
        if k not in seen:
            uniq.append(x)
            seen.add(k)
    return uniq[:60]


# ------------------------------------------------------------
# Additional metadata extraction (scope, submission, contacts, etc.)
# ------------------------------------------------------------

def extract_reference_number(text: str) -> Optional[str]:
    # common tender ID patterns
    patterns = [
        r"(?i)\b(?:no\.?|num(?:éro)?|réf(?:érence)?|reference|solicitation)\s*[:#]?\s*([A-Za-z0-9][A-Za-z0-9\-_/]{3,})\b",
        r"(?i)\b(?:ao|rfp|rfq|rfi|aoi)\s*[:#]?\s*([A-Za-z0-9][A-Za-z0-9\-_/]{3,})\b",
    ]
    for pat in patterns:
        m = re.search(pat, text or "")
        if m:
            return m.group(1)[:80]
    return None


def extract_contract_type(text: str) -> Optional[str]:
    low = (text or "").lower()
    candidates = [
        ("appel d'offres", "Appel d’offres"),
        ("appel d’intérêt", "Appel d’intérêt"),
        ("demande de propositions", "Demande de propositions"),
        ("demande de prix", "Demande de prix"),
        ("rfi", "RFI"),
        ("rfp", "RFP"),
        ("rfq", "RFQ"),
        ("aoi", "AOI"),
    ]
    for needle, label in candidates:
        if needle in low:
            return label
    return None


def extract_language(text: str) -> Optional[str]:
    snip = _find_best_snippet(text, ["langue", "language", "français", "anglais", "bilingual", "bilingue"], 200, 900)
    if not snip:
        return None
    low = snip.lower()
    if "biling" in low or ("français" in low and "anglais" in low):
        return "FR/EN"
    if "français" in low or "francais" in low:
        return "FR"
    if "anglais" in low or "english" in low:
        return "EN"
    return None


def extract_submission_info(text: str) -> Dict[str, Any]:
    snip = _find_best_snippet(text, ["soumission", "submission", "dépôt", "depot", "remise", "transmettre", "submit"], 250, 1500) or ""
    emails = _extract_emails(snip)
    platforms = []
    low = snip.lower()
    for p in ["seao", "canadabuys", "merx", "bonfire", "biddingo", "sap ariba", "ariba", "tenders", "achat"]:
        if p in low:
            platforms.append(p.upper() if p.isalpha() else p)

    # crude address capture: lines with postal codes
    addr = None
    for line in (snip.splitlines() if snip else []):
        if re.search(r"(?i)\b[abceghj-nprstvxy]\d[abceghj-nprstvxy][ -]?\d[abceghj-nprstvxy]\d\b", line):
            addr = _clip(line, 200)
            break

    return {
        "snippet": snip or None,
        "platforms": list(dict.fromkeys(platforms))[:6],
        "emails": emails,
        "address_hint": addr,
    }


def extract_scope_summary(text: str) -> Optional[str]:
    # Try to pick the "Objet / Mandat / Contexte / Besoins" section
    snip = _find_best_snippet(text, ["objet", "description", "besoin", "mandat", "contexte", "scope of work", "statement of work"], 200, 2200)
    return snip


def extract_security_terms(text: str) -> List[str]:
    anchors = [
        "cote de sécurité", "cote de securite", "attestation de sécurité", "attestation de securite",
        "habilitat", "secret", "protégé", "protege", "security clearance", "reliability", "protected",
    ]
    snip = _find_best_snippet(text, anchors, 200, 1200)
    if not snip:
        return []
    lines = [l.strip() for l in snip.splitlines() if l.strip()]
    out = []
    for l in lines:
        ll = l.lower()
        if any(a in ll for a in anchors) and 20 <= len(l) <= 260:
            out.append(_normalize_spaces(l))
    # de-dup
    uniq, seen = [], set()
    for x in out:
        k = x.lower()
        if k not in seen:
            uniq.append(x)
            seen.add(k)
    return uniq[:15]


def extract_key_contacts(text: str) -> Dict[str, Any]:
    snip = _find_best_snippet(text, ["contact", "renseignements", "information", "procurement", "acheteur", "buyer"], 250, 1600) or ""
    emails = _extract_emails(snip)
    phones = _extract_phones(snip)
    return {
        "emails": emails,
        "phones": phones,
        "snippet": snip or None,
    }


def _keyword_overlap_score(hay: str, needles: List[str]) -> int:
    low = (hay or "").lower()
    return sum(1 for n in needles if n.lower() in low)


def compute_fit_score(fields: Dict[str, Any], user: Optional[AuthenticatedUser]) -> Dict[str, Any]:
    """
    Very simple "go/no-go" helper.
    We do NOT decide for the user, we give a signal and reasons.
    """
    score = 50
    reasons: List[str] = []

    title = str(fields.get("title") or "")
    scope = str(fields.get("scope_summary") or "")
    combined = f"{title}\n{scope}"

    it_hits = _keyword_overlap_score(combined, ["ti", "informat", "service", "logiciel", "erp", "oracle", "servicenow", "ia", "ai", "data", "integration"])
    if it_hits >= 2:
        score += 12
        reasons.append("Le contenu semble aligné avec des mots-clés TI/ERP/IA.")
    elif it_hits == 1:
        score += 5
        reasons.append("Le contenu contient au moins un mot-clé TI/ERP/IA.")

    if fields.get("mandatory_requirements"):
        score += 3
        reasons.append("Des exigences obligatoires ont été détectées (à valider).")

    # Deadline proximity
    kd = fields.get("key_dates") or {}
    closing = kd.get("closing_date") or fields.get("closing_date")
    if closing:
        try:
            y, m, d = [int(x) for x in str(closing).split("-")]
            delta = (date(y, m, d) - date.today()).days
            if delta < 0:
                score -= 30
                reasons.append("Date de clôture passée (vérifier si addenda / prolongation).")
            elif delta <= 7:
                score -= 10
                reasons.append("Délai court avant clôture (risque opérationnel).")
        except Exception:
            pass

    # Security clearance can be a blocker
    sec = fields.get("security_terms") or []
    if sec:
        score -= 8
        reasons.append("Mention de sécurité/habilitation détectée (peut être bloquant selon votre statut).")

    # User specialty (if available)
    if user:
        specialty = (user.profile.main_specialty or "") + " " + (user.profile.activity_type or "")
        overlap = _keyword_overlap_score(specialty, ["erp", "ia", "oracle", "servicenow", "data", "integration", "it"])
        if overlap:
            score += 5
            reasons.append("Ton profil utilisateur (spécialité) semble pertinent pour ce type d’AO.")

    score = max(0, min(score, 100))
    verdict = "GO" if score >= 60 else ("MAYBE" if score >= 45 else "NO-GO")
    return {"score": score, "verdict": verdict, "reasons": reasons[:8]}


def build_compliance_checklist(fields: Dict[str, Any]) -> List[str]:
    checklist: List[str] = []
    kd = fields.get("key_dates") or {}
    if kd.get("questions_deadline"):
        checklist.append(f"Soumettre les questions avant: {kd.get('questions_deadline')}")

    if kd.get("closing_date") or fields.get("closing_date"):
        checklist.append(f"Déposer la soumission avant: {kd.get('closing_date') or fields.get('closing_date')}")

    # Always include baseline docs
    checklist.extend([
        "Lettre de présentation (résumé exécutif + engagement)",
        "Méthodologie / approche + plan de travail",
        "Équipe proposée (CV + rôles + disponibilités)",
        "Références / projets similaires",
        "Offre financière (structure de prix / taux / hypothèses)",
        "Formulaires et attestations demandés (si applicables)",
    ])

    # If we detected docs, add them explicitly
    docs = fields.get("submission_documents") or []
    for d in docs[:15]:
        checklist.append(f"Doc requis: {d}")

    # Mandatory requirements
    for m in (fields.get("mandatory_requirements") or [])[:12]:
        checklist.append(f"Obligatoire: {m}")

    # De-dup
    out, seen = [], set()
    for x in checklist:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out[:35]


def build_proposal_outline(fields: Dict[str, Any]) -> List[str]:
    """
    A proposal-ready outline (structure), not the full proposal.
    """
    outline = [
        "1. Résumé exécutif",
        "2. Compréhension du besoin et contexte",
        "3. Portée (scope) et hypothèses",
        "4. Approche/méthodologie (phases, activités, livrables)",
        "5. Gouvernance et gestion de projet (rôles, cadence, qualité)",
        "6. Équipe et compétences (CV, certifications, disponibilité)",
        "7. Planification (jalons, calendrier, dépendances)",
        "8. Livrables (détail + format)",
        "9. Gestion des risques (identification, mitigation)",
        "10. Conformité aux exigences obligatoires (table de conformité)",
        "11. Références et expériences pertinentes",
        "12. Offre financière (structure de prix, hypothèses, limites)",
        "13. Annexes (formulaires, attestations, documents requis)",
    ]
    if fields.get("evaluation_criteria"):
        outline.insert(10, "10. Alignement sur les critères d’évaluation (réponse point par point)")
        outline[10+1] = "11. Conformité aux exigences obligatoires (table de conformité)"
    return outline


def build_risks(fields: Dict[str, Any]) -> List[str]:
    risks: List[str] = []
    if fields.get("security_terms"):
        risks.append("Exigences de sécurité/habilitation: vérifier admissibilité de l’équipe/entreprise.")
    if not (fields.get("estimated_value")):
        risks.append("Budget/valeur non détecté: risque sur l’effort vs. rentabilité; clarifier si possible.")
    if not (fields.get("submission_documents")):
        risks.append("Liste des documents requis non détectée: risque de non-conformité; valider dans le cahier.")
    if not (fields.get("mandatory_requirements")):
        risks.append("Exigences obligatoires non identifiées automatiquement: revoir la section conformité.")
    if (fields.get("key_dates") or {}).get("site_visit_date"):
        risks.append("Visite / réunion: confirmer si obligatoire et prévoir la présence.")
    return risks[:12]


# ------------------------------------------------------------
# DB enrichment (ao.db)
# ------------------------------------------------------------

def _db_get_tender_by_id(tender_id: int) -> Dict[str, Any]:
    if tender_id is None:
        return {}
    if not AO_DB_PATH.exists():
        return {}

    possible_tables = ["tenders", "ao", "ao_tenders", "tender"]
    possible_cols = ["id", "title", "url", "portal_name", "published_at", "buyer", "country", "region"]

    try:
        con = sqlite3.connect(str(AO_DB_PATH))
        con.row_factory = sqlite3.Row
    except Exception:
        return {}

    try:
        cur = con.cursor()
        tables = [r["name"] for r in cur.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()]
        table = next((t for t in possible_tables if t in tables), None)
        if not table:
            return {}

        cols = [r["name"] for r in cur.execute(f"PRAGMA table_info({table});").fetchall()]
        wanted = [c for c in possible_cols if c in cols]
        if "id" not in cols:
            return {}

        sel = ", ".join(wanted) if wanted else "*"
        row = cur.execute(f"SELECT {sel} FROM {table} WHERE id = ?", (tender_id,)).fetchone()
        if not row:
            return {}

        d = dict(row)
        out = {
            "title": d.get("title"),
            "url": d.get("url"),
            "portal_name": d.get("portal_name") or d.get("source"),
            "published_at": d.get("published_at"),
            "buyer": d.get("buyer"),
            "country": d.get("country"),
            "region": d.get("region"),
        }
        return {k: v for k, v in out.items() if v not in (None, "", "null")}
    except Exception:
        return {}
    finally:
        try:
            con.close()
        except Exception:
            pass


# ------------------------------------------------------------
# Summary, confidence, next actions
# ------------------------------------------------------------

def build_summary(text_len: int, extracted_fields: Dict[str, Any]) -> str:
    parts = [f"Texte extrait ({text_len} caractères)."]
    if extracted_fields.get("closing_date"):
        parts.append(f"Clôture: {extracted_fields.get('closing_date')}.")
    if extracted_fields.get("buyer"):
        parts.append(f"Acheteur: {extracted_fields.get('buyer')}.")
    if extracted_fields.get("estimated_value"):
        parts.append(f"Valeur estimée: {extracted_fields.get('estimated_value')}.")
    if extracted_fields.get("reference_number"):
        parts.append(f"Référence: {extracted_fields.get('reference_number')}.")
    if extracted_fields.get("portal_name"):
        parts.append(f"Portail: {extracted_fields.get('portal_name')}.")
    return " ".join(parts).strip()


def compute_confidence(text_len: int, extracted_fields: Dict[str, Any]) -> float:
    score = 0.25
    if text_len > 20_000:
        score += 0.20
    if text_len > 80_000:
        score += 0.15

    if extracted_fields.get("closing_date"):
        score += 0.10
    if extracted_fields.get("buyer"):
        score += 0.10
    if extracted_fields.get("reference_number"):
        score += 0.06
    if extracted_fields.get("scope_summary"):
        score += 0.08

    if extracted_fields.get("mandatory_requirements"):
        score += 0.10
    if extracted_fields.get("deliverables"):
        score += 0.08
    if extracted_fields.get("evaluation_criteria"):
        score += 0.08
    if extracted_fields.get("submission_documents"):
        score += 0.07

    return float(max(0.2, min(score, 0.95)))


def build_next_actions(fields: Dict[str, Any]) -> List[str]:
    actions: List[str] = []

    kd = fields.get("key_dates") or {}
    if kd.get("closing_date"):
        actions.append(f"Valider la date de clôture ({kd.get('closing_date')}) dans le document")
    else:
        actions.append("Extraire les dates clés (clôture / visite / questions)")

    if kd.get("questions_deadline"):
        actions.append(f"Préparer et envoyer les questions avant ({kd.get('questions_deadline')})")

    if not fields.get("submission_documents"):
        actions.append("Identifier la liste des documents à fournir (formulaires, attestations, annexes)")

    if not fields.get("mandatory_requirements"):
        actions.append("Identifier les exigences obligatoires (éliminatoires) et préparer une table de conformité")

    if not fields.get("deliverables"):
        actions.append("Lister les livrables attendus et les transformer en plan de projet / WBS")

    if not fields.get("evaluation_criteria"):
        actions.append("Lister les critères d’évaluation + pondération et répondre point par point")

    if not fields.get("contact_emails"):
        actions.append("Identifier le contact officiel (email/téléphone) pour clarifications")

    out, seen = [], set()
    for a in actions:
        if a not in seen:
            out.append(a)
            seen.add(a)
    return out[:20]


# ------------------------------------------------------------
# HTML Report rendering
# ------------------------------------------------------------

def _html_escape(s: Any) -> str:
    s = "" if s is None else str(s)
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _html_kv(rows: List[Tuple[str, Any]]) -> str:
    out = []
    for k, v in rows:
        vv = "" if v in (None, "", [], {}) else v
        out.append(f'<div class="k">{_html_escape(k)}</div><div class="v">{_html_escape(vv)}</div>')
    return "<div class='kv'>" + "".join(out) + "</div>"


def render_report_html(analysis: Dict[str, Any]) -> str:
    result = analysis.get("result", {}) or {}
    fields = result.get("extracted_fields", {}) or {}

    title = fields.get("title") or analysis.get("db_enriched", {}).get("title") or "Rapport d’analyse AO"
    key_dates = fields.get("key_dates") or {}

    def li(items: List[str]) -> str:
        if not items:
            return '<div class="muted">Non détecté.</div>'
        return "<ul>" + "".join(f"<li>{_html_escape(x)}</li>" for x in items) + "</ul>"

    fit = result.get("fit") or {}
    fit_badge = f"{fit.get('verdict')} ({fit.get('score')})" if fit else "N/A"

    return f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{_html_escape(title)}</title>
  <style>
    :root {{
      --bg: #0b1220;
      --card: rgba(255,255,255,.06);
      --border: rgba(255,255,255,.12);
      --text: rgba(255,255,255,.92);
      --muted: rgba(255,255,255,.65);
      --accent: #7dd3fc;
      --good: #86efac;
      --warn: #fde68a;
      --bad: #fca5a5;
    }}
    *{{box-sizing:border-box}}
    body {{
      margin:0;
      font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial;
      background: radial-gradient(1200px 700px at 20% -10%, rgba(125,211,252,.20), transparent 60%),
                  radial-gradient(900px 600px at 110% 0%, rgba(134,239,172,.12), transparent 55%),
                  var(--bg);
      color: var(--text);
      padding: 28px;
    }}
    .wrap {{ max-width: 1020px; margin: 0 auto; }}
    .top {{
      display:flex; gap:16px; align-items:flex-start; justify-content:space-between; flex-wrap:wrap;
      margin-bottom: 18px;
    }}
    .h1 {{ font-size: 22px; font-weight: 800; letter-spacing: .2px; margin: 0; }}
    .badge {{
      display:inline-flex; align-items:center; gap:10px;
      padding: 8px 12px; border: 1px solid var(--border); border-radius: 999px;
      background: rgba(255,255,255,.04);
      font-size: 13px; color: var(--muted);
    }}
    .grid {{
      display:grid; grid-template-columns: 1fr 1fr;
      gap: 14px;
    }}
    @media (max-width: 980px) {{
      .grid {{ grid-template-columns: 1fr; }}
    }}
    .card {{
      background: var(--card);
      border: 1px solid var(--border);
      border-radius: 16px;
      padding: 14px 14px;
      box-shadow: 0 10px 30px rgba(0,0,0,.25);
      backdrop-filter: blur(10px);
    }}
    .card h2 {{
      margin: 0 0 10px 0;
      font-size: 13px;
      text-transform: uppercase;
      letter-spacing: .12em;
      color: rgba(255,255,255,.80);
    }}
    .kv {{ display:grid; grid-template-columns: 190px 1fr; gap: 8px 14px; }}
    .k {{ color: var(--muted); font-size: 13px; }}
    .v {{ font-size: 13px; }}
    a {{ color: var(--accent); text-decoration: none; }}
    a:hover {{ text-decoration: underline; }}
    ul {{ margin: 8px 0 0 18px; padding: 0; }}
    li {{ margin: 6px 0; color: rgba(255,255,255,.85); font-size: 13px; line-height: 1.35; }}
    .muted {{ color: var(--muted); font-size: 13px; }}
    .bar {{
      height: 8px; background: rgba(255,255,255,.08);
      border: 1px solid var(--border); border-radius: 999px;
      overflow: hidden; margin-top: 8px;
    }}
    .fill {{
      height: 100%;
      width: {int((result.get("confidence") or 0.0) * 100)}%;
      background: linear-gradient(90deg, rgba(125,211,252,.85), rgba(134,239,172,.85));
    }}
    .actions a {{
      display:inline-block; margin-right: 10px; margin-top: 8px;
      padding: 8px 12px; border-radius: 12px;
      border: 1px solid var(--border);
      background: rgba(255,255,255,.04);
      font-size: 13px;
    }}
    .tag {{
      display:inline-block; padding: 4px 10px; border-radius: 999px;
      border: 1px solid var(--border); background: rgba(255,255,255,.04);
      font-size: 12px; color: rgba(255,255,255,.85);
      margin-right: 8px;
    }}
    pre {{
      white-space: pre-wrap;
      background: rgba(0,0,0,.20);
      border: 1px solid var(--border);
      padding: 10px;
      border-radius: 12px;
      color: rgba(255,255,255,.85);
      font-size: 12px;
      line-height: 1.35;
    }}
  </style>
</head>
<body>
<div class="wrap">
  <div class="top">
    <div>
      <h1 class="h1">{_html_escape(title)}</h1>
      <div class="muted" style="margin-top:6px;">Analysis ID: {_html_escape(analysis.get("analysis_id"))} · {_html_escape(analysis.get("created_at"))}</div>
      <div style="margin-top:10px;">
        <span class="tag">FIT: {_html_escape(fit_badge)}</span>
        <span class="tag">Confiance: {_html_escape(result.get("confidence"))}</span>
        <span class="tag">Réf: {_html_escape(fields.get("reference_number"))}</span>
      </div>
    </div>
    <div class="badge">
      <span>Confiance</span>
      <strong style="color: var(--good);">{_html_escape(result.get("confidence"))}</strong>
      <div class="bar" style="width:120px;"><div class="fill"></div></div>
    </div>
  </div>

  <div class="card" style="margin-bottom:14px;">
    <h2>Résumé exécutif</h2>
    <div class="muted">{_html_escape(result.get("summary"))}</div>
    <div class="actions">
      <a href="/api/ai/report/docx/{_html_escape(analysis.get("analysis_id"))}">Télécharger Word (DOCX)</a>
      <a href="/api/ai/report/pdf/{_html_escape(analysis.get("analysis_id"))}">Télécharger PDF</a>
    </div>
    <div class="muted" style="margin-top:10px;">
      <strong>GO/NO-GO:</strong> {_html_escape(fit.get("verdict"))} · {_html_escape(fit.get("score"))}
      <br/>
      {_html_escape("; ".join(fit.get("reasons") or []))}
    </div>
  </div>

  <div class="grid">
    <div class="card">
      <h2>Fiche AO</h2>
      {_html_kv([
        ("Portail", fields.get("portal_name")),
        ("Publication", fields.get("published_at")),
        ("Acheteur", fields.get("buyer")),
        ("Valeur estimée", fields.get("estimated_value")),
        ("Type", fields.get("contract_type")),
        ("Langue", fields.get("language")),
        ("URL", fields.get("url") or ""),
      ])}
    </div>

    <div class="card">
      <h2>Dates clés</h2>
      {_html_kv([
        ("Clôture", key_dates.get("closing_date") or fields.get("closing_date")),
        ("Questions", key_dates.get("questions_deadline")),
        ("Visite / réunion", key_dates.get("site_visit_date")),
        ("Addenda", key_dates.get("addenda_deadline")),
        ("Ouverture", key_dates.get("opening_date")),
      ])}
    </div>

    <div class="card" style="grid-column: 1 / -1;">
      <h2>Portée / Contexte (extrait)</h2>
      <pre>{_html_escape(fields.get("scope_summary") or "Non détecté.")}</pre>
    </div>

    <div class="card">
      <h2>Exigences obligatoires (détectées)</h2>
      {li(fields.get("mandatory_requirements") or [])}
    </div>

    <div class="card">
      <h2>Documents de soumission (détectés)</h2>
      {li(fields.get("submission_documents") or [])}
    </div>

    <div class="card">
      <h2>Livrables</h2>
      {li(fields.get("deliverables") or [])}
      <h2 style="margin-top:14px;">Critères d’évaluation</h2>
      {li(fields.get("evaluation_criteria") or [])}
    </div>

    <div class="card">
      <h2>Contacts (indices)</h2>
      {_html_kv([
        ("Emails", ", ".join(fields.get("contact_emails") or [])),
        ("Téléphones", ", ".join(fields.get("contact_phones") or [])),
        ("Plateformes", ", ".join((fields.get("submission") or {}).get("platforms") or [])),
        ("Adresse (indice)", (fields.get("submission") or {}).get("address_hint")),
      ])}
    </div>

    <div class="card" style="grid-column: 1 / -1;">
      <h2>Checklist conformité (prêt à soumettre)</h2>
      {li(result.get("compliance_checklist") or [])}
    </div>

    <div class="card" style="grid-column: 1 / -1;">
      <h2>Plan de réponse (structure de proposition)</h2>
      {li(result.get("proposal_outline") or [])}
    </div>

    <div class="card" style="grid-column: 1 / -1;">
      <h2>Risques / points d’attention</h2>
      {li(result.get("risks") or [])}
    </div>

    <div class="card" style="grid-column: 1 / -1;">
      <h2>Prochaines actions</h2>
      {li(result.get("next_actions") or [])}
    </div>
  </div>
</div>
</body>
</html>
"""


# ------------------------------------------------------------
# DOCX / PDF generation
# ------------------------------------------------------------

def _analysis_to_plain_sections(analysis: Dict[str, Any]) -> Dict[str, Any]:
    result = analysis.get("result", {}) or {}
    fields = result.get("extracted_fields", {}) or {}
    return {
        "title": fields.get("title") or "Rapport d’analyse AO",
        "summary": result.get("summary"),
        "confidence": result.get("confidence"),
        "fit": result.get("fit") or {},
        "fields": fields,
        "next_actions": result.get("next_actions") or [],
        "compliance_checklist": result.get("compliance_checklist") or [],
        "proposal_outline": result.get("proposal_outline") or [],
        "risks": result.get("risks") or [],
    }


def render_docx_bytes(analysis: Dict[str, Any]) -> bytes:
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise HTTPException(
            status_code=501,
            detail=f"Export DOCX indisponible (python-docx non installé). Installe: pip install python-docx. Détails: {e}",
        )

    s = _analysis_to_plain_sections(analysis)
    fields = s["fields"]
    kd = fields.get("key_dates") or {}
    submission = fields.get("submission") or {}

    doc = Document()
    doc.add_heading(str(s["title"]), level=0)
    doc.add_paragraph(f"Confiance: {s['confidence']}")
    if s.get("fit"):
        doc.add_paragraph(f"GO/NO-GO: {s['fit'].get('verdict')} ({s['fit'].get('score')})")

    doc.add_paragraph(str(s["summary"] or ""))

    doc.add_heading("Fiche AO", level=1)
    for k, label in [
        ("reference_number", "Référence"),
        ("portal_name", "Portail"),
        ("published_at", "Publication"),
        ("buyer", "Acheteur"),
        ("estimated_value", "Valeur estimée"),
        ("contract_type", "Type"),
        ("language", "Langue"),
        ("url", "URL"),
    ]:
        if fields.get(k) not in (None, "", [], {}):
            doc.add_paragraph(f"{label}: {fields.get(k)}")

    doc.add_heading("Dates clés", level=1)
    for k, label in [
        ("closing_date", "Clôture"),
        ("questions_deadline", "Questions"),
        ("site_visit_date", "Visite / réunion"),
        ("addenda_deadline", "Addenda"),
        ("opening_date", "Ouverture"),
    ]:
        v = kd.get(k) if isinstance(kd, dict) else None
        doc.add_paragraph(f"{label}: {v}")

    doc.add_heading("Portée / Contexte (extrait)", level=1)
    doc.add_paragraph(str(fields.get("scope_summary") or "Non détecté."))

    def add_list(title: str, items: List[str]) -> None:
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph("Non détecté.")
            return
        for x in items:
            doc.add_paragraph(str(x), style="List Bullet")

    add_list("Exigences obligatoires (détectées)", fields.get("mandatory_requirements") or [])
    add_list("Documents de soumission (détectés)", fields.get("submission_documents") or [])
    add_list("Livrables", fields.get("deliverables") or [])
    add_list("Critères d’évaluation", fields.get("evaluation_criteria") or [])

    doc.add_heading("Contacts (indices)", level=1)
    if fields.get("contact_emails"):
        doc.add_paragraph(f"Emails: {', '.join(fields.get('contact_emails') or [])}")
    if fields.get("contact_phones"):
        doc.add_paragraph(f"Téléphones: {', '.join(fields.get('contact_phones') or [])}")
    if submission.get("platforms"):
        doc.add_paragraph(f"Plateformes: {', '.join(submission.get('platforms') or [])}")
    if submission.get("address_hint"):
        doc.add_paragraph(f"Adresse (indice): {submission.get('address_hint')}")

    add_list("Checklist conformité (prêt à soumettre)", s.get("compliance_checklist") or [])
    add_list("Plan de réponse (structure de proposition)", s.get("proposal_outline") or [])
    add_list("Risques / points d’attention", s.get("risks") or [])
    add_list("Prochaines actions", s.get("next_actions") or [])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf_bytes(analysis: Dict[str, Any]) -> bytes:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise HTTPException(
            status_code=501,
            detail=f"Export PDF indisponible (reportlab non installé). Installe: pip install reportlab. Détails: {e}",
        )

    s = _analysis_to_plain_sections(analysis)
    fields = s["fields"]
    kd = fields.get("key_dates") or {}
    submission = fields.get("submission") or {}

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER

    def wrap_lines(txt: str, max_chars: int = 105) -> List[str]:
        txt = _normalize_spaces(str(txt or ""))
        if not txt:
            return [""]
        lines = []
        while txt:
            lines.append(txt[:max_chars])
            txt = txt[max_chars:]
        return lines

    def draw_para(txt: str, y: float, bold: bool = False) -> float:
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 10 if not bold else 11)
        for line in wrap_lines(txt):
            c.drawString(0.75 * inch, y, line)
            y -= 12
            if y < 0.9 * inch:
                c.showPage()
                y = height - 0.9 * inch
                c.setFont("Helvetica", 10)
        return y

    y = height - 0.9 * inch
    y = draw_para(s["title"], y, bold=True)
    y = draw_para(f"Confiance: {s['confidence']} | GO/NO-GO: {s.get('fit', {}).get('verdict')} ({s.get('fit', {}).get('score')})", y)
    y = draw_para(s["summary"] or "", y)
    y -= 6

    def section(title: str) -> None:
        nonlocal y
        y -= 4
        y = draw_para(title, y, bold=True)

    def kv_line(label: str, value: Any) -> None:
        nonlocal y
        if value in (None, "", [], {}):
            return
        y = draw_para(f"{label}: {value}", y)

    section("Fiche AO")
    kv_line("Référence", fields.get("reference_number"))
    kv_line("Portail", fields.get("portal_name"))
    kv_line("Publication", fields.get("published_at"))
    kv_line("Acheteur", fields.get("buyer"))
    kv_line("Valeur estimée", fields.get("estimated_value"))
    kv_line("Type", fields.get("contract_type"))
    kv_line("Langue", fields.get("language"))
    kv_line("URL", fields.get("url"))
    y -= 6

    section("Dates clés")
    kv_line("Clôture", kd.get("closing_date") or fields.get("closing_date"))
    kv_line("Questions", kd.get("questions_deadline"))
    kv_line("Visite / réunion", kd.get("site_visit_date"))
    kv_line("Addenda", kd.get("addenda_deadline"))
    kv_line("Ouverture", kd.get("opening_date"))
    y -= 6

    section("Portée / Contexte (extrait)")
    y = draw_para(fields.get("scope_summary") or "Non détecté.", y)
    y -= 6

    def list_block(title: str, items: List[str], limit: int = 18) -> None:
        nonlocal y
        section(title)
        if not items:
            y = draw_para("Non détecté.", y)
            return
        for it in items[:limit]:
            y = draw_para(f"- {it}", y)

    list_block("Exigences obligatoires (détectées)", fields.get("mandatory_requirements") or [], 14)
    list_block("Documents de soumission (détectés)", fields.get("submission_documents") or [], 14)
    list_block("Livrables", fields.get("deliverables") or [], 14)
    list_block("Critères d’évaluation", fields.get("evaluation_criteria") or [], 14)

    section("Contacts (indices)")
    kv_line("Emails", ", ".join(fields.get("contact_emails") or []))
    kv_line("Téléphones", ", ".join(fields.get("contact_phones") or []))
    kv_line("Plateformes", ", ".join(submission.get("platforms") or []))
    kv_line("Adresse (indice)", submission.get("address_hint"))
    y -= 6

    list_block("Checklist conformité (prêt à soumettre)", s.get("compliance_checklist") or [], 18)
    list_block("Plan de réponse (structure de proposition)", s.get("proposal_outline") or [], 18)
    list_block("Risques / points d’attention", s.get("risks") or [], 12)
    list_block("Prochaines actions", s.get("next_actions") or [], 12)

    c.showPage()
    c.save()
    return buf.getvalue()


# ------------------------------------------------------------
# API routes
# ------------------------------------------------------------

@router.post("/analyze")
async def analyze_ao(
    user: AuthenticatedUser = Depends(get_current_user),
    tender_id: Optional[int] = Form(default=None),
    notes: Optional[str] = Form(default=None),
    files: List[UploadFile] = File(...),
) -> Dict[str, Any]:
    if not files:
        raise HTTPException(status_code=422, detail="Aucun fichier reçu.")

    combined_parts: List[str] = []
    file_meta: List[Dict[str, Any]] = []

    for f in files:
        content = await f.read()
        file_meta.append({"filename": f.filename, "content_type": f.content_type, "size_bytes": len(content)})

        if (f.content_type or "").lower() != "application/pdf" and not (f.filename or "").lower().endswith(".pdf"):
            raise HTTPException(status_code=422, detail=f"Fichier non-PDF: {f.filename}")

        txt = extract_text_from_pdf(content)
        combined_parts.append(f"\n\n===== FILE: {f.filename} =====\n\n{txt}")

    combined_text = "\n".join(combined_parts).strip()
    text_len = len(combined_text)

    # PDF parsing
    buyer = _pick_buyer(combined_text)
    estimated_value = _pick_estimated_value(combined_text)
    key_dates = extract_key_dates(combined_text)

    mandatory_requirements = extract_mandatory_requirements(combined_text)
    deliverables = extract_deliverables(combined_text)
    evaluation_criteria = extract_evaluation_criteria(combined_text)
    submission_documents = extract_submission_documents(combined_text)

    # Extra metadata
    reference_number = extract_reference_number(combined_text)
    contract_type = extract_contract_type(combined_text)
    language = extract_language(combined_text)
    scope_summary = extract_scope_summary(combined_text)
    contacts = extract_key_contacts(combined_text)
    submission = extract_submission_info(combined_text)
    security_terms = extract_security_terms(combined_text)

    db_enriched = _db_get_tender_by_id(int(tender_id)) if tender_id is not None else {}

    extracted_fields: Dict[str, Any] = {
        "tender_id": tender_id,
        "reference_number": reference_number,
        "closing_date": key_dates.get("closing_date"),
        "buyer": buyer or db_enriched.get("buyer"),
        "estimated_value": estimated_value,
        "title": db_enriched.get("title"),
        "url": db_enriched.get("url"),
        "portal_name": db_enriched.get("portal_name"),
        "published_at": db_enriched.get("published_at"),
        "country": db_enriched.get("country"),
        "region": db_enriched.get("region"),
        "contract_type": contract_type,
        "language": language,
        "key_dates": key_dates,
        "scope_summary": scope_summary,
        "mandatory_requirements": mandatory_requirements,
        "submission_documents": submission_documents,
        "deliverables": deliverables,
        "evaluation_criteria": evaluation_criteria,
        "submission": submission,
        "contact_emails": contacts.get("emails") or [],
        "contact_phones": contacts.get("phones") or [],
        "security_terms": security_terms,
    }
    extracted_fields = {k: v for k, v in extracted_fields.items() if v not in ("", None)}

    # Build higher-level analysis helpers
    fit = compute_fit_score(extracted_fields, user=user)
    compliance_checklist = build_compliance_checklist(extracted_fields)
    proposal_outline = build_proposal_outline(extracted_fields)
    risks = build_risks(extracted_fields)
    next_actions = build_next_actions(extracted_fields)

    summary = build_summary(text_len=text_len, extracted_fields=extracted_fields)
    confidence = compute_confidence(text_len=text_len, extracted_fields=extracted_fields)

    analysis_id = f"ana_{uuid.uuid4().hex[:12]}"
    created_at = _utc_iso()

    payload: Dict[str, Any] = {
        "status": "ok",
        "analysis_id": analysis_id,
        "created_at": created_at,
        "inputs": {"tender_id": tender_id, "notes": notes, "file_count": len(files), "files": file_meta},
        "result": {
            "summary": summary,
            "extracted_fields": extracted_fields,
            "fit": fit,
            "compliance_checklist": compliance_checklist,
            "proposal_outline": proposal_outline,
            "risks": risks,
            "next_actions": next_actions,
            "confidence": round(confidence, 2),
            "debug": {"text_chars": text_len, "text_sample": combined_text[:1200]},
        },
        "user": user.profile.model_dump(),
        "report_urls": {
            "html": f"/api/ai/report/html/{analysis_id}",
            "docx": f"/api/ai/report/docx/{analysis_id}",
            "pdf": f"/api/ai/report/pdf/{analysis_id}",
        },
        "db_enriched": db_enriched,
    }

    ANALYSES[analysis_id] = payload
    return payload


@router.get("/report/html/{analysis_id}", response_class=HTMLResponse)
def report_html(analysis_id: str, user: AuthenticatedUser = Depends(get_current_user)) -> HTMLResponse:
    analysis = ANALYSES.get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis introuvable (serveur redémarré ?). Relance l'analyse.")
    html = render_report_html(analysis)
    return HTMLResponse(content=html, status_code=200)


@router.get("/report/docx/{analysis_id}")
def report_docx(analysis_id: str, user: AuthenticatedUser = Depends(get_current_user)) -> StreamingResponse:
    analysis = ANALYSES.get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis introuvable (serveur redémarré ?). Relance l'analyse.")
    try:
        b = render_docx_bytes(analysis)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("DOCX generation failed")
        raise HTTPException(status_code=500, detail=f"Erreur génération DOCX: {e}")

    title = _safe_filename((analysis.get("result", {}).get("extracted_fields", {}) or {}).get("title") or analysis_id)
    filename = f"{title}__{analysis_id}.docx"
    return StreamingResponse(
        io.BytesIO(b),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/report/pdf/{analysis_id}")
def report_pdf(analysis_id: str, user: AuthenticatedUser = Depends(get_current_user)) -> StreamingResponse:
    analysis = ANALYSES.get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis introuvable (serveur redémarré ?). Relance l'analyse.")
    try:
        b = render_pdf_bytes(analysis)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("PDF generation failed")
        raise HTTPException(status_code=500, detail=f"Erreur génération PDF: {e}")

    title = _safe_filename((analysis.get("result", {}).get("extracted_fields", {}) or {}).get("title") or analysis_id)
    filename = f"{title}__{analysis_id}.pdf"
    return StreamingResponse(
        io.BytesIO(b),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
